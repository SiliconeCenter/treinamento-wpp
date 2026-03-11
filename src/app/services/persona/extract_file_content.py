import pathlib
import shutil
from pathlib import Path

import docx
from fastapi import UploadFile
from PyPDF2 import PdfReader


def read_file(file: UploadFile):
    # verifica a extensão do arquivo:
    upload_dir = Path("data/")
    upload_dir.mkdir(parents=True, exist_ok=True)  # Cria a pasta se não existir

    dest = upload_dir / file.filename
    with dest.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    path = str(dest)

    tipo = pathlib.Path(path).suffix

    match tipo:
        case ".txt":
            with open(path, "r", encoding="utf-8") as f:
                conteudo = f.read()
                return conteudo

        case ".pdf":
            reader = PdfReader(path)
            texto = ""

            for pagina in reader.pages:
                texto += pagina.extract_text() + "\n"

            return texto

        case ".doc" | ".docx":
            try:
                doc = docx.Document(path)
                texto = ""

                # Extrai texto dos parágrafos
                for paragrafo in doc.paragraphs:
                    if paragrafo.text.strip():  # Ignora linhas vazias
                        texto += paragrafo.text + "\n"

                # Extrai texto das tabelas (se houver)
                for tabela in doc.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            if celula.text.strip():
                                texto += celula.text + " | "
                        texto += "\n"

                return texto if texto.strip() else "Arquivo Word vazio"

            except Exception as e:
                return f"Erro ao ler arquivo Word: {str(e)}"
        case _:
            return "O tipo do arquivo não é suportado!"
